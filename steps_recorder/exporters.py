"""Xuất báo cáo: HTML tự chứa, Markdown (ảnh assets), DOCX.

Các hàm nhận StepsRecorder làm tham số đầu (không import để tránh vòng lặp).
"""
import base64
import datetime as dt
import io
import logging
import os
from typing import List, Optional

log = logging.getLogger("steps_recorder")


def _now() -> str:
    return dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _esc(text: str) -> str:
    return (str(text).replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;").replace('"', "&quot;"))


def export_html(recorder, path: str, title: Optional[str] = None,
                include_toc: bool = True):
    title = title or recorder.report_title or "Bản ghi các bước"
    rows = []
    # Nhóm bước theo section (giữ thứ tự xuất hiện) cho mục lục lồng nhau
    section_order: List[str] = []
    section_steps: dict = {}  # section_key -> list[Step]
    has_any_section = any((s.section or "").strip() for s in recorder.steps)
    for s in recorder.steps:
        sec = (s.section or "").strip()
        if not sec and has_any_section:
            sec = "Các bước khác"
        key = sec or ""  # rỗng = không chia phần
        if key not in section_steps:
            section_steps[key] = []
            section_order.append(key)
        section_steps[key].append(s)

    prev_section = object()  # sentinel
    for s in recorder.steps:
        if s.images:
            imgs_html = (
                '<figure class="shots">'
                + "".join(
                    f'<img src="data:image/png;base64,{b}" '
                    f'alt="Minh họa bước {s.index}: {_esc(s.action)}">'
                    for b in s.images)
                + "</figure>")
        else:
            imgs_html = ""
        if s.description:
            desc_body = "<br>".join(
                _esc(line) if line.strip() else "<br>"
                for line in s.description.splitlines())
            desc_html = f'<div class="desc">{desc_body}</div>'
        else:
            desc_html = ""
        win_html = (
            f'<p class="window">Ứng dụng / cửa sổ: <em>{_esc(s.window)}</em></p>'
            if s.window and s.window != "(không xác định)" else "")
        sec_name = (s.section or "").strip()
        if not sec_name and has_any_section:
            sec_name = "Các bước khác"
        heading_html = ""
        if sec_name and sec_name != prev_section:
            sec_idx = (section_order.index(sec_name) + 1
                       if sec_name in section_order else 0)
            heading_html = (
                f'<header class="part-head" id="sec-{sec_idx}">'
                f'<h2 class="part-title">{_esc(sec_name)}</h2></header>')
            prev_section = sec_name
        rows.append(f"""
        {heading_html}
        <section class="step" id="step-{s.index}" data-step="{s.index}"
                 data-section="{_esc(sec_name)}">
          <div class="head">
            <span class="badge">Bước {s.index}</span>
            <h2 class="action">{_esc(s.action)}</h2>
            <span class="time">{_esc(s.timestamp)}</span>
          </div>
          {desc_html}
          {win_html}
          {imgs_html}
        </section>""")

    # Mục lục: lồng bước trong từng phần (nếu có section)
    nav_blocks = []
    if has_any_section:
        for i, sec_key in enumerate(section_order, start=1):
            steps_in = section_steps[sec_key]
            children = "".join(
                f'<li><a class="nav-link" href="#step-{st.index}" '
                f'data-step="{st.index}">'
                f'<span class="nav-n">{st.index}</span>'
                f'<span class="nav-t">{_esc(st.action)}</span></a></li>'
                for st in steps_in)
            nav_blocks.append(
                f'<li class="nav-part">'
                f'<a class="nav-part-link" href="#sec-{i}" data-step="sec-{i}">'
                f'<span class="nav-part-t">Phần {i}: {_esc(sec_key)}</span></a>'
                f'<ol class="nav-children">{children}</ol></li>')
    else:
        for s in recorder.steps:
            nav_blocks.append(
                f'<li><a class="nav-link" href="#step-{s.index}" data-step="{s.index}">'
                f'<span class="nav-n">{s.index}</span>'
                f'<span class="nav-t">{_esc(s.action)}</span></a></li>')

    summary_html = ""
    summary_nav = ""
    if recorder.report_summary:
        sum_body = "<br>".join(
            _esc(line) if line.strip() else "<br>"
            for line in recorder.report_summary.splitlines())
        summary_html = (
            '<section class="summary" id="intro"><h2>Giới thiệu</h2>'
            f'<p>{sum_body}</p></section>'
        )
        summary_nav = (
            '<li><a class="nav-link" href="#intro" data-step="intro">'
            '<span class="nav-n">i</span>'
            '<span class="nav-t">Giới thiệu</span></a></li>')

    n_parts = len(section_order) if has_any_section else 0
    meta_parts = (
        f" · {n_parts} phần" if n_parts else "")
    sidebar_html = ""
    if include_toc and recorder.steps:
        sidebar_html = f"""
  <aside class="sidebar" id="sidebar" aria-label="Mục lục các bước">
<div class="sidebar-head">
  <div class="sidebar-title">Mục lục</div>
  <div class="sidebar-meta">{len(recorder.steps)} bước{meta_parts}</div>
</div>
<nav class="side-nav">
  <ol>
    {summary_nav}
    {''.join(nav_blocks)}
  </ol>
</nav>
  </aside>"""

    html = f"""<!DOCTYPE html>
<html lang="vi"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{_esc(title)}</title>
<style>
  :root {{
--brand:#0b5cab; --brand-soft:#e8f1fb; --brand-active:#0a4d8c;
--text:#1f2328; --muted:#57606a; --line:#d0d7de; --bg:#f6f8fa;
--card:#fff; --sidebar-w:280px;
  }}
  * {{ box-sizing:border-box; }}
  html {{ scroll-behavior:smooth; }}
  body {{ font-family:"Segoe UI", system-ui, -apple-system, sans-serif; margin:0;
     background:var(--bg); color:var(--text); line-height:1.55; }}
  .topbar {{ background:linear-gradient(135deg,#0b5cab 0%,#1f6feb 100%);
         color:#fff; padding:20px 28px; position:sticky; top:0; z-index:20;
         box-shadow:0 1px 0 rgba(0,0,0,.08); }}
  .topbar .doc-type {{ display:inline-block; font-size:11px; letter-spacing:.06em;
                   text-transform:uppercase; opacity:.9; margin-bottom:6px;
                   border:1px solid rgba(255,255,255,.35); border-radius:999px;
                   padding:2px 10px; }}
  .topbar h1 {{ margin:0; font-size:22px; font-weight:700; line-height:1.3; }}
  .topbar .meta {{ margin:8px 0 0; opacity:.88; font-size:13px; }}
  .layout {{ display:flex; align-items:flex-start; min-height:calc(100vh - 110px); }}
  .sidebar {{
position:sticky; top:110px; align-self:flex-start;
width:var(--sidebar-w); flex:0 0 var(--sidebar-w);
height:calc(100vh - 110px); overflow:auto;
background:#fff; border-right:1px solid var(--line);
padding:0 0 24px;
  }}
  .sidebar-head {{ padding:16px 16px 10px; border-bottom:1px solid var(--line);
               position:sticky; top:0; background:#fff; z-index:1; }}
  .sidebar-title {{ font-size:12px; font-weight:700; color:var(--brand);
                text-transform:uppercase; letter-spacing:.05em; }}
  .sidebar-meta {{ font-size:12px; color:var(--muted); margin-top:2px; }}
  .side-nav ol {{ list-style:none; margin:8px 0 0; padding:0 8px; }}
  .side-nav li {{ margin:2px 0; }}
  .nav-link, .nav-part-link {{
display:flex; gap:10px; align-items:flex-start;
text-decoration:none; color:var(--text); padding:8px 10px;
border-radius:8px; font-size:13px; line-height:1.35;
border-left:3px solid transparent;
  }}
  .nav-link:hover, .nav-part-link:hover {{ background:var(--brand-soft); color:var(--brand); }}
  .nav-link.active, .nav-part-link.active {{
background:var(--brand-soft); color:var(--brand-active);
border-left-color:var(--brand); font-weight:600;
  }}
  .nav-n, .nav-part-n {{
flex:0 0 1.6em; min-width:1.6em; height:1.6em; line-height:1.6em;
text-align:center; font-size:11px; font-weight:700;
background:#eef2f6; color:var(--muted); border-radius:999px;
  }}
  .nav-link.active .nav-n, .nav-part-link.active .nav-part-n {{
background:var(--brand); color:#fff;
  }}
  .nav-t, .nav-part-t {{ flex:1; word-break:break-word; }}
  .nav-part {{ margin:8px 0 4px; }}
  .nav-part-link {{
font-weight:700; font-size:12.5px; color:var(--brand-active);
text-transform:none; padding:7px 10px 5px;
  }}
  .nav-part-n {{
background:var(--brand-soft); color:var(--brand);
  }}
  .nav-children {{
list-style:none; margin:2px 0 6px 0; padding:0 0 0 6px;
border-left:2px solid #e8eef5;
  }}
  .nav-children .nav-link {{ padding:6px 8px 6px 10px; font-size:12.5px; }}
  .content {{ flex:1; min-width:0; max-width:920px; padding:24px 22px 48px; }}
  .summary, .step {{
background:var(--card); border:1px solid var(--line); border-radius:12px;
padding:18px 22px; margin-bottom:20px;
  }}
  .summary {{ border-left:5px solid var(--brand); }}
  .summary h2 {{ margin:0 0 10px; font-size:15px; color:var(--brand);
             text-transform:uppercase; letter-spacing:.04em; }}
  .summary p {{ margin:0; color:#24292f; }}
  .part-head {{
margin:28px 0 14px; padding:0 2px 10px;
border-bottom:2px solid var(--brand);
scroll-margin-top:120px;
  }}
  .part-head:first-child {{ margin-top:4px; }}
  .part-title {{
margin:0; font-size:17px; font-weight:700; color:var(--brand-active);
letter-spacing:.01em;
  }}
  .step {{ scroll-margin-top:128px; }}
  .head {{ display:grid; grid-template-columns:auto 1fr auto; gap:10px 14px;
       align-items:start; border-bottom:1px solid #eaeef2;
       padding-bottom:12px; margin-bottom:12px; }}
  .badge {{ display:inline-block; background:var(--brand-soft); color:var(--brand);
        font-weight:700; font-size:12px; padding:4px 10px; border-radius:999px;
        white-space:nowrap; margin-top:2px; }}
  .action {{ margin:0; font-size:18px; font-weight:650; line-height:1.35;
         color:#0d1117; grid-column:2; }}
  .time {{ font-size:12px; color:var(--muted); white-space:nowrap; margin-top:4px; }}
  .desc {{ font-size:15px; margin:0 0 10px; color:#24292f; }}
  .window {{ font-size:12px; color:var(--muted); margin:0 0 10px; }}
  .shots {{ margin:12px 0 0; padding:0; }}
  .shots img {{ max-width:100%; border:1px solid var(--line); border-radius:8px;
            display:block; margin-top:10px; box-shadow:0 1px 2px rgba(0,0,0,.04); }}
  footer {{ max-width:920px; padding:0 22px 36px; font-size:12px;
        color:var(--muted); text-align:center; }}
  .menu-toggle {{
display:none; position:fixed; bottom:18px; left:18px; z-index:30;
background:var(--brand); color:#fff; border:0; border-radius:999px;
padding:10px 16px; font-size:13px; font-weight:600; cursor:pointer;
box-shadow:0 4px 14px rgba(11,92,171,.35);
  }}
  @media (max-width:900px) {{
.layout {{ display:block; }}
.sidebar {{
  position:fixed; left:0; top:0; bottom:0; z-index:40;
  width:min(86vw, 300px); height:100vh; transform:translateX(-105%);
  transition:transform .2s ease; box-shadow:4px 0 20px rgba(0,0,0,.12);
}}
body.nav-open .sidebar {{ transform:translateX(0); }}
.menu-toggle {{ display:inline-flex; align-items:center; gap:6px; }}
.content {{ max-width:none; padding:18px 14px 72px; }}
.step {{ scroll-margin-top:100px; }}
  }}
  @media print {{
body {{ background:#fff; }}
.topbar {{ position:static; background:#0b5cab !important;
           -webkit-print-color-adjust:exact; print-color-adjust:exact; }}
.sidebar, .menu-toggle {{ display:none !important; }}
.layout {{ display:block; }}
.content {{ max-width:none; padding:12px 0; }}
.step, .summary {{ break-inside:avoid; box-shadow:none; }}
  }}
</style></head>
<body>
  <header class="topbar">
<div class="doc-type">Hướng dẫn sử dụng</div>
<h1>{_esc(title)}</h1>
<p class="meta">Ban hành: {_now()} · Tổng số bước: {len(recorder.steps)}</p>
  </header>
  <div class="layout">
{sidebar_html}
<div class="content">
  {summary_html}
  {''.join(rows)}
  <footer>Tài liệu được biên soạn từ bản ghi thao tác · {len(recorder.steps)} bước</footer>
</div>
  </div>
  <button type="button" class="menu-toggle" id="menuToggle" aria-label="Mở mục lục">☰ Mục lục</button>
<script>
(function () {{
  var links = Array.prototype.slice.call(
document.querySelectorAll(".nav-link, .nav-part-link"));
  var sections = Array.prototype.slice.call(
document.querySelectorAll(".step, .summary, .part-head"));
  var body = document.body;
  var btn = document.getElementById("menuToggle");
  var sidebar = document.getElementById("sidebar");

  function setActive(id) {{
links.forEach(function (a) {{
  var href = a.getAttribute("href") || "";
  a.classList.toggle("active", href === "#" + id);
}});
  }}

  links.forEach(function (a) {{
a.addEventListener("click", function () {{
  body.classList.remove("nav-open");
  var href = a.getAttribute("href") || "";
  if (href.charAt(0) === "#") setActive(href.slice(1));
}});
  }});

  if (btn) {{
btn.addEventListener("click", function () {{
  body.classList.toggle("nav-open");
}});
  }}
  document.addEventListener("click", function (e) {{
if (!body.classList.contains("nav-open")) return;
if (sidebar && sidebar.contains(e.target)) return;
if (btn && btn.contains(e.target)) return;
body.classList.remove("nav-open");
  }});

  if ("IntersectionObserver" in window && sections.length) {{
var io = new IntersectionObserver(function (entries) {{
  var visible = entries
    .filter(function (en) {{ return en.isIntersecting; }})
    .sort(function (a, b) {{ return b.intersectionRatio - a.intersectionRatio; }});
  if (visible[0] && visible[0].target && visible[0].target.id) {{
    setActive(visible[0].target.id);
    var act = document.querySelector(".nav-link.active");
    if (act && sidebar) {{
      var r = act.getBoundingClientRect();
      var sr = sidebar.getBoundingClientRect();
      if (r.top < sr.top + 40 || r.bottom > sr.bottom - 20) {{
        act.scrollIntoView({{ block: "nearest" }});
      }}
    }}
  }}
}}, {{ rootMargin: "-30% 0px -55% 0px", threshold: [0.1, 0.4, 0.7] }});
sections.forEach(function (s) {{ io.observe(s); }});
  }}
  if (location.hash) setActive(location.hash.slice(1));
  else if (document.getElementById("intro")) setActive("intro");
  else if (sections[0]) setActive(sections[0].id);
}})();
</script>
</body></html>"""

    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    return path

def export_markdown(recorder, path: str, title: Optional[str] = None) -> str:
    """Xuất Markdown; ảnh giải mã ra thư mục <tên file>_assets bên cạnh
    (data-URI không hiển thị được trên GitHub)."""
    title = title or recorder.report_title or "Bản ghi các bước"
    base = os.path.splitext(os.path.basename(path))[0]
    assets_dir = os.path.join(os.path.dirname(path) or ".", f"{base}_assets")

    has_any_section = any((s.section or "").strip() for s in recorder.steps)
    lines: List[str] = [f"# {title}", ""]
    if recorder.report_summary.strip():
        lines += [recorder.report_summary.strip(), ""]

    wrote_assets = False
    prev_section = None
    for s in recorder.steps:
        sec = (s.section or "").strip()
        if not sec and has_any_section:
            sec = "Các bước khác"
        if sec and sec != prev_section:
            lines += [f"## {sec}", ""]
            prev_section = sec
        lines.append(f"### Bước {s.index}: {s.action}")
        if s.window and s.window != "(không xác định)":
            lines.append(f"*Ứng dụng / cửa sổ: {s.window}*")
        lines.append("")
        if s.description.strip():
            lines += [s.description.strip(), ""]
        for j, b64 in enumerate(s.images):
            if not wrote_assets:
                os.makedirs(assets_dir, exist_ok=True)
                wrote_assets = True
            img_name = f"step_{s.index}_{j + 1}.png"
            with open(os.path.join(assets_dir, img_name), "wb") as imf:
                imf.write(base64.b64decode(b64))
            lines.append(
                f"![Minh họa bước {s.index}]({base}_assets/{img_name})")
        if s.images:
            lines.append("")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines).rstrip() + "\n")
    return path




def export_docx(recorder, path: str, title: Optional[str] = None) -> str:
    """Xuất tài liệu Word (.docx): tiêu đề, tóm tắt, phần, bước, ảnh nhúng."""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as e:
        raise RuntimeError(
            "Thiếu thư viện 'python-docx'. Cài bằng: pip install python-docx") from e
    from PIL import Image

    title = title or recorder.report_title or "Bản ghi các bước"
    doc = Document()
    doc.add_heading(title, level=0)
    if recorder.report_summary.strip():
        doc.add_paragraph(recorder.report_summary.strip())

    has_any_section = any((s.section or "").strip() for s in recorder.steps)
    prev_section = None
    for s in recorder.steps:
        sec = (s.section or "").strip()
        if not sec and has_any_section:
            sec = "Các bước khác"
        if sec and sec != prev_section:
            doc.add_heading(sec, level=1)
            prev_section = sec
        doc.add_heading(f"Bước {s.index}: {s.action}", level=2)
        if s.window and s.window != "(không xác định)":
            p = doc.add_paragraph()
            p.add_run(f"Ứng dụng / cửa sổ: {s.window}").italic = True
        if s.description.strip():
            doc.add_paragraph(s.description.strip())
        for b64 in s.images:
            try:
                raw = base64.b64decode(b64)
                # Ảnh rộng thu về bề ngang trang (~6 inch); ảnh nhỏ giữ nguyên
                px_w = Image.open(io.BytesIO(raw)).width
                kwargs = {"width": Inches(6.0)} if px_w > 600 else {}
                doc.add_picture(io.BytesIO(raw), **kwargs)
            except Exception:
                log.warning("Bỏ qua một ảnh không đọc được ở bước %s", s.index)

    if not path.lower().endswith(".docx"):
        path += ".docx"
    doc.save(path)
    return path
